import io
import zipfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

employees_db = {
    "Alex_Rivera": {
        "role": "Senior Backend Engineer",
        "manager": "Jordan Smith",
        "months": [
            {
                "date": "October 2025",
                "summary": "Alex continues to be a high-performing member of the engineering team. This month, velocity was incredible, and the API schema migration was completed three days ahead of schedule. However, I’ve noticed the junior devs are struggling to understand the new code because the documentation pages are empty. We need to balance speed with maintainability.",
                "kpis": [
                    ("System Uptime", "99.9%", "99.95%", "Exceeded"),
                    ("Story Points", "30", "45", "Exceeded"),
                    ("Documentation Score", "4/5", "2/5", "Needs Improvement")
                ],
                "competencies": [
                    "**Technical Proficiency:** (Expert) The migration strategy was flawless.",
                    "**Mentorship:** (Developing) Junior devs are blocked by lack of docs."
                ],
                "strengths": [
                    "Coding velocity is currently best-in-class for the team.",
                    "Deep knowledge of legacy codebase allowed for quick refactoring."
                ],
                "development": [
                    "Documentation habits need immediate attention.",
                    "Tendency to work in a silo during 'crunch' periods."
                ],
                "growth": [
                    "Dedicate the last Friday of the sprint solely to updating the internal Wiki.",
                    "Pair program with a junior dev on the next API endpoint."
                ]
            },
            {
                "date": "November 2025",
                "summary": "This month highlighted Alex’s value as a Senior Engineer. Identifying that authentication vulnerability saved us from a potential breach. The 'Crisis Mode' work was excellent, but it meant the documentation goal from October was missed again. Alex is working late too often; we need to watch for burnout.",
                "kpis": [
                    ("System Uptime", "99.9%", "99.99%", "Exceeded"),
                    ("Critical Bugs Fixed", "0", "3", "Exceeded"),
                    ("Docs Updated", "5 pages", "1 page", "Missed")
                ],
                "competencies": [
                    "**Problem Solving:** (Expert) Caught a critical auth bug before prod.",
                    "**Reliability:** (High) Was available during the weekend incident."
                ],
                "strengths": [
                    "Crisis management and calmness under pressure.",
                    "Security-first mindset."
                ],
                "development": [
                    "Work-life balance; logged 60 hours this week.",
                    "Documentation backlog is growing."
                ],
                "growth": [
                    "Take two comp days off post-incident.",
                    "No new feature work next sprint; focus strictly on tech debt."
                ]
            },
            {
                "date": "December 2025",
                "summary": "Technically, Alex crushed the Q4 goals. All features shipped. However, the Frontend team expressed frustration that Alex went 'radio silent' for three days during the integration phase. While the code worked, the lack of communication caused anxiety for other departments.",
                "kpis": [
                    ("Feature Completion", "100%", "100%", "Met"),
                    ("Communication Score", "4/5", "2/5", "Needs Focus"),
                    ("Peer Reviews", "10", "15", "Exceeded")
                ],
                "competencies": [
                    "**Execution:** (High) Delivered 'Project Nexus' on time.",
                    "**Collaboration:** (Low) Failed to update ticket status for 72 hours."
                ],
                "strengths": [
                    "Reliable delivery of complex features.",
                    "Code quality is exceptionally high (very few QA returns)."
                ],
                "development": [
                    "Communication timing; updates must happen *before* blockers arise.",
                    "Cross-functional empathy."
                ],
                "growth": [
                    "Mandatory attendance at the daily cross-functional standup.",
                    "Setup automated status reminders in Slack."
                ]
            },
            {
                "date": "January 2026",
                "summary": "A fantastic reset. Alex took the feedback about communication to heart. The documentation is now the 'gold standard' for the team. Alex also spent significant time pair-programming with our two new hires, which has accelerated their onboarding.",
                "kpis": [
                    ("Uptime", "99.9%", "100%", "Exceeded"),
                    ("Mentorship Hours", "5", "15", "Exceeded"),
                    ("Docs Updated", "5 pages", "12 pages", "Exceeded")
                ],
                "competencies": [
                    "**Communication:** (Improved) Proactive updates in Slack all month.",
                    "**Leadership:** (Emerging) Took lead on onboarding new hires."
                ],
                "strengths": [
                    "Adaptability to feedback.",
                    "Teaching ability; new hires praised Alex's patience."
                ],
                "development": [
                    "Coding velocity dipped slightly due to mentorship (expected).",
                    "Delegation; still hesitant to hand off complex tasks."
                ],
                "growth": [
                    "Lead the 'Architecture Brown Bag' session next month.",
                    "Identify one module to fully hand over to a mid-level engineer."
                ]
            },
            {
                "date": "February 2026",
                "summary": "Alex has successfully pivoted from 'executing' to 'leading.' The proposal for the new RAG (Retrieval-Augmented Generation) system was well-researched and approved by the CTO. Alex is effectively delegating lower-level tasks to the juniors he mentored in Jan.",
                "kpis": [
                    ("R&D Prototypes", "1", "1", "Met"),
                    ("System Uptime", "99.9%", "99.98%", "Met"),
                    ("Leadership Score", "3/5", "5/5", "Exceeded")
                ],
                "competencies": [
                    "**Strategic Thinking:** (High) RAG proposal aligned with company 2026 goals.",
                    "**Delegation:** (Improved) Successfully handed off the payment module."
                ],
                "strengths": [
                    "Visionary thinking regarding AI integration.",
                    "Trustworthiness with senior leadership."
                ],
                "development": [
                    "Public speaking; presentation to CTO was slightly nervous.",
                    "Time estimation for R&D projects."
                ],
                "growth": [
                    "Enroll in 'Engineering Management 101' track.",
                    "Present the RAG roadmap to the entire company at All-Hands."
                ]
            }
        ]
    },
    "Sarah_Jenkins": {
        "role": "Junior Marketing Associate",
        "manager": "Elena Rodriguez",
        "months": [
            {
                "date": "October 2025",
                "summary": "Sarah has settled in well. She brings great energy and visual creativity. However, we had two incidents this month where posts went out with typos. At this stage, attention to detail is just as important as creativity. We need to implement a proofing process.",
                "kpis": [
                    ("Social Posts", "15", "15", "Met"),
                    ("Typos/Errors", "0", "2", "Needs Improvement"),
                    ("Engagement Rate", "2.0%", "1.8%", "Near Target")
                ],
                "competencies": [
                    "**Creativity:** (High) Visuals are stunning.",
                    "**Attention to Detail:** (Low) Proofreading needs work."
                ],
                "strengths": [
                    "Graphic design skills (Canva/Photoshop).",
                    "Enthusiasm for brand voice."
                ],
                "development": [
                    "Written accuracy (grammar/spelling).",
                    "Understanding the approval workflow."
                ],
                "growth": [
                    "Install Grammarly browser extension.",
                    "All posts must be reviewed by a peer before scheduling."
                ]
            },
            {
                "date": "November 2025",
                "summary": "Much better quality control this month. Zero errors in published content. Sarah is now struggling slightly with the analytics side of the role; she can create the content but has trouble explaining *why* a certain post performed well during our weekly reviews.",
                "kpis": [
                    ("Social Posts", "15", "18", "Exceeded"),
                    ("Analytics Reports", "4", "4", "Met"),
                    ("Error Rate", "0%", "0%", "Met")
                ],
                "competencies": [
                    "**Data Literacy:** (Developing) Struggling to interpret engagement metrics.",
                    "**Consistency:** (High) Never missed a posting deadline."
                ],
                "strengths": [
                    "Reliability; deadlines are always met.",
                    "Improved attention to detail."
                ],
                "development": [
                    "Data analysis; moving beyond 'vanity metrics'.",
                    "Strategic planning."
                ],
                "growth": [
                    "Shadow the Data Analyst for one hour/week.",
                    "Complete Google Analytics Academy basics course."
                ]
            },
            {
                "date": "December 2025",
                "summary": "Breakthrough month! Sarah identified a rising TikTok trend and quickly pivoted our content calendar to match it. The resulting video went viral (50k+ views), giving us our best brand awareness month of the year. She proved she understands our audience intuitively.",
                "kpis": [
                    ("Viral Views", "5k", "52k", "Exceeded"),
                    ("New Followers", "100", "850", "Exceeded"),
                    ("Campaign ROI", "3x", "10x", "Exceeded")
                ],
                "competencies": [
                    "**Innovation:** (Expert) Spotted trend before competitors.",
                    "**Agility:** (High) Pivoted strategy in 24 hours."
                ],
                "strengths": [
                    "Trend spotting.",
                    "Video editing speed."
                ],
                "development": [
                    "Time management; the viral post took 3 days, delaying other tasks.",
                    "Managing inbox overflow from new followers."
                ],
                "growth": [
                    "Create a 'Trend Response' protocol so we can move fast without dropping other balls.",
                    "Lead a team teach-in on video editing."
                ]
            },
            {
                "date": "January 2026",
                "summary": "Post-viral stabilization. Engagement has normalized. Sarah is working on converting those new followers into leads. She is now much more comfortable with data, correctly identifying that our new audience prefers educational content over sales content.",
                "kpis": [
                    ("Lead Conversion", "1.0%", "0.8%", "Near Target"),
                    ("Posts", "20", "20", "Met"),
                    ("Report Quality", "3/5", "4/5", "Improved")
                ],
                "competencies": [
                    "**Strategic Insight:** (Improved) Connected content types to conversion rates.",
                    "**Resilience:** (High) Handled the drop in views post-viral well."
                ],
                "strengths": [
                    "Data-backed decision making.",
                    "Audience engagement."
                ],
                "development": [
                    "Copywriting for conversion (sales focus).",
                    "Email marketing integration."
                ],
                "growth": [
                    "A/B test call-to-action buttons on next 5 posts.",
                    "Collaborate with Sales team to understand customer pain points."
                ]
            },
            {
                "date": "February 2026",
                "summary": "Sarah is evolving from a content creator to a holistic marketer. She successfully launched the Q1 newsletter with minimal supervision. Her presentation of the monthly metrics to the VP of Marketing was clear, concise, and highlighted actionable insights.",
                "kpis": [
                    ("Newsletter Open Rate", "20%", "28%", "Exceeded"),
                    ("Lead Gen", "50", "65", "Exceeded"),
                    ("Presentation Score", "3/5", "5/5", "Exceeded")
                ],
                "competencies": [
                    "**Communication:** (High) Executive presentation was excellent.",
                    "**Autonomy:** (High) Managed newsletter end-to-end."
                ],
                "strengths": [
                    "Full-stack marketing (Design + Copy + Data).",
                    "Confidence in presentations."
                ],
                "development": [
                    "Long-term campaign planning (quarterly vs monthly).",
                    "Budget management."
                ],
                "growth": [
                    "Given ownership of the Q2 'Spring Fling' campaign budget ($5k).",
                    "Mentoring the new marketing intern."
                ]
            }
        ]
    },
    "David_Chen": {
        "role": "HR Manager",
        "manager": "Marcus Thorne",
        "months": [
            {
                "date": "October 2025",
                "summary": "David is the heart of the office, and his conflict resolution skills are unmatched. However, the administrative side of the role is lagging. We are seeing delays in offer letters being sent out because David is resisting the move to the new HRIS system, preferring manual spreadsheets.",
                "kpis": [
                    ("Time to Hire", "30 days", "45 days", "Missed"),
                    ("Employee Satisfaction", "4.0", "4.8", "Exceeded"),
                    ("System Adoption", "100%", "40%", "Critical Issue")
                ],
                "competencies": [
                    "**Empathy:** (Expert) Resolved a difficult dispute between Sales/Support.",
                    "**Digital Fluency:** (Low) Refusal to use the new dashboard."
                ],
                "strengths": [
                    "Interpersonal relationships.",
                    "Cultural advocacy."
                ],
                "development": [
                    "Technical adaptability.",
                    "Process efficiency."
                ],
                "growth": [
                    "Mandatory 1:1 training with the IT implementation lead.",
                    "Goal: Send 100% of offer letters via the new system in Nov."
                ]
            },
            {
                "date": "November 2025",
                "summary": "Significant improvement on the technical front. David bit the bullet and learned the new system. While slower than usual, he processed all compliance documents correctly. He also led a fantastic 'Mental Health Awareness' week that received high praise from staff.",
                "kpis": [
                    ("Compliance Audit", "100%", "100%", "Met"),
                    ("System Usage", "100%", "90%", "Met"),
                    ("Event Feedback", "4.0", "4.9", "Exceeded")
                ],
                "competencies": [
                    "**Adaptability:** (Improved) Overcame resistance to software.",
                    "**Culture Building:** (High) Mental Health week was a hit."
                ],
                "strengths": [
                    "Event planning and execution.",
                    "Trust among employees."
                ],
                "development": [
                    "Speed of processing in the new system.",
                    "Data reporting."
                ],
                "growth": [
                    "Create a custom dashboard for recruitment stats.",
                    "Automate the birthday/anniversary email triggers."
                ]
            },
            {
                "date": "December 2025",
                "summary": "David organized the best Holiday Party we've had in years, staying perfectly within budget. However, during the end-of-year bonus calculations, a spreadsheet error nearly caused overpayments. We caught it in time, but this reinforces the need to move *everything* off Excel.",
                "kpis": [
                    ("Budget Adherence", "100%", "100%", "Met"),
                    ("Party Attendance", "80%", "98%", "Exceeded"),
                    ("Payroll Accuracy", "100%", "95%", "Near Miss")
                ],
                "competencies": [
                    "**Planning:** (High) complex event management.",
                    "**Attention to Detail:** (Mixed) Excel formula error was dangerous."
                ],
                "strengths": [
                    "Budget management.",
                    "Morale boosting."
                ],
                "development": [
                    "Data integrity validation.",
                    "Moving away from legacy manual processes."
                ],
                "growth": [
                    "Work with Finance to audit the bonus calculation process.",
                    "Retire the 'Master HR Spreadsheet' permanently by Jan 1."
                ]
            },
            {
                "date": "January 2026",
                "summary": "A very strong start to the year. The 'Master Spreadsheet' is gone, and David is now using the HRIS for all payroll inputs. The efficiency gain is visible—payroll took 2 days instead of 5. He is now focused on the annual performance review cycle.",
                "kpis": [
                    ("Payroll Processing", "5 days", "2 days", "Exceeded"),
                    ("Review Completion", "90%", "85%", "On Track"),
                    ("Hiring pipeline", "10 cands", "12 cands", "Met")
                ],
                "competencies": [
                    "**Efficiency:** (High) Digital transformation is paying off.",
                    "**Process Management:** (High) Smooth rollout of review cycles."
                ],
                "strengths": [
                    "Operational efficiency (newfound).",
                    "Coaching managers on how to write reviews."
                ],
                "development": [
                    "Strategic workforce planning.",
                    "Analyzing retention trends."
                ],
                "growth": [
                    "Prepare a 'Retention Risk' report for the Exec team.",
                    "Analyze exit interview data from 2025."
                ]
            },
            {
                "date": "February 2026",
                "summary": "David is now operating as a strategic partner rather than just an administrator. His presentation on 'Reducing Sales Turnover' was backed by data and offered concrete solutions. He has fully bridged the gap between his high EQ and the necessary operational rigour.",
                "kpis": [
                    ("Retention Rate", "95%", "98%", "Exceeded"),
                    ("Strategic Projects", "1", "1", "Met"),
                    ("Manager NPS", "8", "9", "Exceeded")
                ],
                "competencies": [
                    "**Strategic HR:** (High) Moved from reactive to proactive.",
                    "**Influence:** (High) Convinced Sales VP to change commission structure."
                ],
                "strengths": [
                    "Data-driven storytelling.",
                    "Executive presence."
                ],
                "development": [
                    "Public speaking polish (filler words).",
                    "Scalability of new retention programs."
                ],
                "growth": [
                    "Lead the quarterly All-Hands HR update.",
                    "Design the Q2 internship program."
                ]
            }
        ]
    }
}

# ==========================================
# 2. DOCUMENT GENERATOR FUNCTION
# ==========================================

def create_rich_docx(employee_name, data):
    doc = Document()
    
    # --- STYLES ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- HEADER ---
    header = doc.add_heading(level=0)
    run = header.add_run(f"Performance Review: {data['date']}")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- EMPLOYEE INFO TABLE ---
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Employee & Role
    cell1 = table.cell(0, 0)
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.add_run("Employee: ").bold = True
    p.add_run(employee_name.replace("_", " "))
    
    cell2 = table.cell(0, 1)
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.add_run("Role: ").bold = True
    p.add_run(employees_db[employee_name]["role"])
    
    # Add Manager Row
    row = table.add_row()
    cell1 = row.cells[0]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.add_run("Manager: ").bold = True
    p.add_run(employees_db[employee_name]["manager"])
    
    cell2 = row.cells[1]
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.add_run("Review Cycle: ").bold = True
    p.add_run("Monthly")

    doc.add_paragraph() # Spacer

    # --- 1. EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(data['summary'])

    # --- 2. KPIS (TABLE FORMAT) ---
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = 'Light Shading Accent 1'
    
    # Header Row
    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Target"
    hdr_cells[2].text = "Actual"
    hdr_cells[3].text = "Status"
    
    # Data Rows
    for metric, target, actual, status in data['kpis']:
        row_cells = kpi_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = target
        row_cells[2].text = actual
        row_cells[3].text = status

    doc.add_paragraph()

    # --- 3. COMPETENCY EVALUATION ---
    doc.add_heading('3. Competency Evaluation', level=1)
    for comp in data['competencies']:
        # Split bold part from rest if possible
        p = doc.add_paragraph(style='List Bullet')
        if "**" in comp:
            parts = comp.split("**")
            # parts[1] is the bold text, parts[2] is the rest
            p.add_run(parts[1]).bold = True
            p.add_run(parts[2])
        else:
            p.text = comp

    # --- 4. STRENGTHS ---
    doc.add_heading('4. Strengths & Achievements', level=1)
    for item in data['strengths']:
        doc.add_paragraph(item, style='List Bullet')

    # --- 5. AREAS FOR DEVELOPMENT ---
    doc.add_heading('5. Areas for Development', level=1)
    for item in data['development']:
        doc.add_paragraph(item, style='List Bullet')

    # --- 6. GROWTH PLAN ---
    doc.add_heading('6. Growth Plan & Next Steps', level=1)
    for item in data['growth']:
        doc.add_paragraph(item, style='List Bullet')

    # --- SIGNATURES ---
    doc.add_paragraph()
    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.allow_autofit = True
    
    c1 = sig_table.cell(0,0)
    c1.text = "_________________________\nManager Signature"
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c2 = sig_table.cell(0,1)
    c2.text = "_________________________\nEmployee Signature"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==========================================
# 3. ZIP GENERATOR
# ==========================================

zip_buffer = io.BytesIO()

with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
    for emp_name, emp_data in employees_db.items():
        for i, month_data in enumerate(emp_data["months"]):
            # Create filename: Alex_Rivera_2025-10_Oct.docx
            # Adding index to ensure order if viewed in file explorer
            idx = str(i+1).zfill(2)
            month_short = month_data['date'].split(" ")[0][:3]
            year = month_data['date'].split(" ")[1]
            filename = f"{emp_name}_{year}-{idx}_{month_short}.docx"
            
            # Generate Doc
            docx_stream = create_rich_docx(emp_name, month_data)
            
            # Add to Zip
            zip_file.writestr(filename, docx_stream.getvalue())

zip_buffer.seek(0)

# Save the final zip file
with open("Detailed_Performance_Reports.zip", "wb") as f:
    f.write(zip_buffer.getvalue())

print("Success! Created 'Detailed_Performance_Reports.zip' with 15 detailed documents.")