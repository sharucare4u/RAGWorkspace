import io
import zipfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT  # Fixed import

# ==========================================
# 1. THE DATA (Junior & Entry Level IT)
# ==========================================

employees_db = {
    "Priyesh_Patel": {
        "role": "Junior DevOps Engineer",
        "manager": "Samira Khan",
        "months": [
            {
                "date": "October 2025",
                "summary": "Priyesh has joined the team with excellent Python scripting skills. However, he is still learning our specific CI/CD pipelines. We had a significant incident this month where a 'force push' to the main branch broke the build server for 4 hours. We need to focus on safety protocols over speed.",
                "kpis": [
                    ("Pipeline Uptime", "99.9%", "99.0%", "Missed"),
                    ("Tickets Closed", "10", "15", "Exceeded"),
                    ("Incident Caused", "0", "1", "Critical Issue")
                ],
                "competencies": [
                    "**Scripting:** (High) Wrote a script to clean up old logs in record time.",
                    "**Process Adherence:** (Low) Bypassed checks to push code faster."
                ],
                "strengths": [
                    "Automation mindset.",
                    "Fast learner of new syntax."
                ],
                "development": [
                    "Understanding 'Production Readiness'.",
                    "Git branching strategy compliance."
                ],
                "growth": [
                    "Complete the internal 'Safe Deployment' certification.",
                    "Revoke direct push access to 'Main' branch temporarily."
                ]
            },
            {
                "date": "November 2025",
                "summary": "Much better adherence to protocols this month. Priyesh successfully containerized the legacy payment application using Docker. He asked for help before merging code, which shows maturity. He is still quiet in team meetings and needs to voice his blockers earlier.",
                "kpis": [
                    ("Docker Conversions", "1 App", "1 App", "Met"),
                    ("Incidents Caused", "0", "0", "Met"),
                    ("Communication", "3/5", "2/5", "Needs Focus")
                ],
                "competencies": [
                    "**Containerization:** (Developing) Good grasp of Dockerfiles.",
                    "**Collaboration:** (Low) Struggled silently for 2 days on a networking bug."
                ],
                "strengths": [
                    "Technical curiosity.",
                    "Follows instructions precisely now."
                ],
                "development": [
                    "Asking for help sooner (Timeboxing).",
                    "Participating in daily standups."
                ],
                "growth": [
                    "Present the Docker solution at the next team demo.",
                    "Timebox issues to 1 hour before asking a Senior."
                ]
            },
            {
                "date": "December 2025",
                "summary": "Priyesh helped the team survive the 'End of Year' freeze. He volunteered to monitor the alerts dashboard during the holidays. While he didn't solve complex alerts, his quick escalation to the Seniors reduced our Mean Time to Resolution (MTTR).",
                "kpis": [
                    ("Alert Response Time", "15 mins", "5 mins", "Exceeded"),
                    ("Escalation Accuracy", "90%", "100%", "Exceeded"),
                    ("Tasks Completed", "12", "12", "Met")
                ],
                "competencies": [
                    "**Monitoring:** (Developing) Learning to read Grafana dashboards.",
                    "**Reliability:** (High) Showed up on time for all holiday shifts."
                ],
                "strengths": [
                    "Dependability.",
                    "Alert triage."
                ],
                "development": [
                    "Troubleshooting independent of Seniors.",
                    "Cloud infrastructure (AWS) basics."
                ],
                "growth": [
                    "Complete 'AWS Cloud Practitioner' study guide.",
                    "Shadow a Senior during a live incident."
                ]
            },
            {
                "date": "January 2026",
                "summary": "Priyesh is starting to pay for himself. He identified a set of 'Orphaned Snapshots' in our AWS account and wrote a script to delete them, saving the company $400/month. This initiative was unprompted and highly appreciated by the Director.",
                "kpis": [
                    ("Cost Savings", "$0", "$400/mo", "Exceeded"),
                    ("Script Errors", "0", "0", "Met"),
                    ("AWS Cert Status", "In Progress", "In Progress", "On Track")
                ],
                "competencies": [
                    "**Initiative:** (High) Found and fixed a cost leak.",
                    "**Cloud Fluency:** (Improved) Navigating AWS console confidently."
                ],
                "strengths": [
                    "Cost optimization.",
                    "Proactive problem finding."
                ],
                "development": [
                    "Documentation of his own scripts.",
                    "Terraform (Infrastructure as Code)."
                ],
                "growth": [
                    "Refactor the cost-saving script into a Lambda function.",
                    "Start the Terraform basic training module."
                ]
            },
            {
                "date": "February 2026",
                "summary": "A solid month. Priyesh handled his first minor deployment solo without breaking anything. He is becoming a reliable junior engineer. His documentation has improved, though his commit messages are still a bit vague.",
                "kpis": [
                    ("Deployments", "2", "2", "Met"),
                    ("Success Rate", "100%", "100%", "Met"),
                    ("Docs Written", "1", "1", "Met")
                ],
                "competencies": [
                    "**Deployment:** (Competent) Can handle standard rollouts.",
                    "**Communication:** (Improved) Updates Jira tickets daily."
                ],
                "strengths": [
                    "Steady growth curve.",
                    "Safe hands (no longer risky)."
                ],
                "development": [
                    "Git commit message quality.",
                    "Debugging complex networking issues."
                ],
                "growth": [
                    "Lead the deployment of the 'User Profile' service next sprint.",
                    "Read 'Conventional Commits' guide."
                ]
            }
        ]
    },
    "Maya_Johnson": {
        "role": "IT Support Specialist (L1)",
        "manager": "David Ross",
        "months": [
            {
                "date": "October 2025",
                "summary": "Maya has joined the Helpdesk with a fantastic attitude. The users love her empathy. However, her technical resolution time is double the team average. She spends too long 'chatting' with users instead of fixing the issue. We need to work on efficiency.",
                "kpis": [
                    ("Tickets Closed/Day", "20", "12", "Missed"),
                    ("CSAT Score", "4.5", "5.0", "Exceeded"),
                    ("Avg Handle Time", "10 mins", "22 mins", "Needs Improvement")
                ],
                "competencies": [
                    "**Customer Service:** (Expert) Users feel heard and valued.",
                    "**Technical Diagnostics:** (Novice) Struggles to find root cause."
                ],
                "strengths": [
                    "Empathy and patience.",
                    "De-escalating angry users."
                ],
                "development": [
                    "Speed/Efficiency.",
                    "Knowledge of the ticketing system shortcuts."
                ],
                "growth": [
                    "Shadow the Senior L1 for 2 hours a day.",
                    "Learn the 'Quick Macros' in ServiceNow."
                ]
            },
            {
                "date": "November 2025",
                "summary": "Improvement seen. Maya discovered the 'Password Reset' macro and her speed doubled on those tickets. She still struggles with hardware troubleshooting (printer/laptop configs), often escalating tickets to L2 that she should be able to solve.",
                "kpis": [
                    ("Tickets Closed/Day", "20", "18", "Near Target"),
                    ("Escalation Rate", "10%", "25%", "Too High"),
                    ("CSAT Score", "4.8", "4.9", "Exceeded")
                ],
                "competencies": [
                    "**Tool Usage:** (Improved) Using macros effectively.",
                    "**Hardware Knowledge:** (Low) Unsure how to re-image a laptop."
                ],
                "strengths": [
                    "Software troubleshooting (Office 365).",
                    "Team morale."
                ],
                "development": [
                    "Hardware diagnostics.",
                    "Reducing unnecessary escalations."
                ],
                "growth": [
                    "Complete the 'Laptop Re-imaging' workshop.",
                    "attempt to solve 3 hardware tickets before escalating."
                ]
            },
            {
                "date": "December 2025",
                "summary": "The 'Holiday VPN Crisis'. When the VPN server flaked out, Maya fielded 50 calls in one day. She remained calm while users were panicking. She didn't fix the server (L3 job), but she kept the workforce calm. Her stress management is excellent.",
                "kpis": [
                    ("Call Volume", "30", "55", "Exceeded"),
                    ("User Complaints", "0", "0", "Met"),
                    ("Attendance", "100%", "100%", "Met")
                ],
                "competencies": [
                    "**Crisis Management:** (High) Calm under fire.",
                    "**Communication:** (High) Explained the outage clearly to non-techs."
                ],
                "strengths": [
                    "Resilience.",
                    "Clear communication."
                ],
                "development": [
                    "Deep technical understanding of VPNs.",
                    "Multi-tasking during high volume."
                ],
                "growth": [
                    "Read the network topology documentation.",
                    "Take a comp day for the holiday effort."
                ]
            },
            {
                "date": "January 2026",
                "summary": "Maya noticed that we get 10 tickets a week about 'How to add a signature in Outlook'. She wrote a PDF guide with screenshots and uploaded it to the Knowledge Base. Ticket volume for that issue dropped to zero. This is the proactive mindset we want.",
                "kpis": [
                    ("KB Articles Created", "1", "1", "Met"),
                    ("Tickets Closed/Day", "20", "24", "Exceeded"),
                    ("First Call Resolution", "70%", "82%", "Exceeded")
                ],
                "competencies": [
                    "**Documentation:** (High) Created a clear, user-friendly guide.",
                    "**Proactivity:** (High) Solved a root cause, not just symptoms."
                ],
                "strengths": [
                    "Identifying trends.",
                    "Written communication."
                ],
                "development": [
                    "Mac OS troubleshooting (we are 20% Mac shop).",
                    "Active Directory administration."
                ],
                "growth": [
                    "Get access to Active Directory (Read-Only) for learning.",
                    "Shadow the Mac Support specialist."
                ]
            },
            {
                "date": "February 2026",
                "summary": "Maya is now operating at a 'Senior L1' level. She handles new hire onboarding (laptop setup + account creation) independently. Her technical skills have caught up to her customer service skills. I am recommending her for the 'Tier 2 Preparation' track.",
                "kpis": [
                    ("Onboarding Setups", "5", "5", "Met"),
                    ("CSAT Score", "4.8", "5.0", "Exceeded"),
                    ("Ticket Quality", "4/5", "5/5", "Exceeded")
                ],
                "competencies": [
                    "**Account Management:** (Competent) Active Directory user creation.",
                    "**Mentorship:** (Emerging) Helped the new intern with phone scripts."
                ],
                "strengths": [
                    "Full-stack Helpdesk (Hardware + Software + People).",
                    "Reliability."
                ],
                "development": [
                    "PowerShell scripting basics.",
                    "Networking (DNS/DHCP)."
                ],
                "growth": [
                    "Enroll in CompTIA Network+ course.",
                    "Begin taking 10% of Tier 2 tickets under supervision."
                ]
            }
        ]
    }
}

# ==========================================
# 2. DOCUMENT GENERATOR FUNCTION
# ==========================================

def create_rich_docx(employee_name, data):
    doc = Document()
    
    # --- STYLES ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- HEADER ---
    header = doc.add_heading(level=0)
    run = header.add_run(f"Performance Review: {data['date']}")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- EMPLOYEE INFO TABLE ---
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Employee & Role
    cell1 = table.cell(0, 0)
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.add_run("Employee: ").bold = True
    p.add_run(employee_name.replace("_", " "))
    
    cell2 = table.cell(0, 1)
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.add_run("Role: ").bold = True
    p.add_run(employees_db[employee_name]["role"])
    
    # Add Manager Row
    row = table.add_row()
    cell1 = row.cells[0]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.add_run("Manager: ").bold = True
    p.add_run(employees_db[employee_name]["manager"])
    
    cell2 = row.cells[1]
    cell2.text = ""
    p = cell2.paragraphs[0]
    p.add_run("Review Cycle: ").bold = True
    p.add_run("Monthly")

    doc.add_paragraph() # Spacer

    # --- 1. EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(data['summary'])

    # --- 2. KPIS (TABLE FORMAT) ---
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = 'Light Shading Accent 1'
    
    # Header Row
    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Target"
    hdr_cells[2].text = "Actual"
    hdr_cells[3].text = "Status"
    
    # Data Rows
    for metric, target, actual, status in data['kpis']:
        row_cells = kpi_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = target
        row_cells[2].text = actual
        row_cells[3].text = status

    doc.add_paragraph()

    # --- 3. COMPETENCY EVALUATION ---
    doc.add_heading('3. Competency Evaluation', level=1)
    for comp in data['competencies']:
        p = doc.add_paragraph(style='List Bullet')
        if "**" in comp:
            parts = comp.split("**")
            p.add_run(parts[1]).bold = True
            p.add_run(parts[2])
        else:
            p.text = comp

    # --- 4. STRENGTHS ---
    doc.add_heading('4. Strengths & Achievements', level=1)
    for item in data['strengths']:
        doc.add_paragraph(item, style='List Bullet')

    # --- 5. AREAS FOR DEVELOPMENT ---
    doc.add_heading('5. Areas for Development', level=1)
    for item in data['development']:
        doc.add_paragraph(item, style='List Bullet')

    # --- 6. GROWTH PLAN ---
    doc.add_heading('6. Growth Plan & Next Steps', level=1)
    for item in data['growth']:
        doc.add_paragraph(item, style='List Bullet')

    # --- SIGNATURES ---
    doc.add_paragraph()
    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.allow_autofit = True
    
    c1 = sig_table.cell(0,0)
    c1.text = "_________________________\nManager Signature"
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c2 = sig_table.cell(0,1)
    c2.text = "_________________________\nEmployee Signature"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==========================================
# 3. ZIP GENERATOR
# ==========================================

zip_buffer = io.BytesIO()

with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
    for emp_name, emp_data in employees_db.items():
        for i, month_data in enumerate(emp_data["months"]):
            idx = str(i+1).zfill(2)
            month_short = month_data['date'].split(" ")[0][:3]
            year = month_data['date'].split(" ")[1]
            filename = f"{emp_name}_{year}-{idx}_{month_short}.docx"
            
            docx_stream = create_rich_docx(emp_name, month_data)
            zip_file.writestr(filename, docx_stream.getvalue())

zip_buffer.seek(0)

# Save the final zip file
with open("IT_Junior_Entry_Reports.zip", "wb") as f:
    f.write(zip_buffer.getvalue())

print("Success! Created 'IT_Junior_Entry_Reports.zip' with 10 detailed documents.")